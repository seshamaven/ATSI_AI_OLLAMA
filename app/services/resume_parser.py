"""Service for parsing resumes and extracting text from files."""
import os
import re
import html
import subprocess
import shutil
import tempfile
import zipfile
from io import BytesIO
from typing import Optional
from docx import Document
import PyPDF2

from app.utils.logging import get_logger
from app.utils.safe_logger import safe_extra
from app.utils.cleaning import normalize_text

logger = get_logger(__name__)

# Try to import OCR libraries
try:
    import pytesseract
    from PIL import Image
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    np = None
    logger.warning("OCR libraries (pytesseract, PIL, opencv) not available. Image-based resumes cannot be processed without OCR.")

# Try to import pdf2image for scanned PDF OCR
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logger.debug("pdf2image not available. Scanned PDF OCR will not be available.")

# Try to import PyMuPDF (fitz) as fallback for PDF image extraction (doesn't require poppler)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None
    logger.debug("PyMuPDF (fitz) not available. Will use pdf2image only for PDF OCR.")

# Try to import HTML parsing libraries
try:
    from bs4 import BeautifulSoup
    HTML_PARSING_AVAILABLE = True
except ImportError:
    HTML_PARSING_AVAILABLE = False
    logger.warning("BeautifulSoup not available. HTML-based resumes cannot be processed without it.")

# Try to import Apache Tika for .doc file support (PRIMARY METHOD - Currently Working)
try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    logger.warning("Apache Tika not available. .doc files cannot be processed without Tika.")

# Try to import olefile for .doc file support (fallback method)
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    logger.debug("olefile not available. Will use Tika only for .doc files.")

# Check for antiword command-line tool
ANTIWORD_AVAILABLE = shutil.which("antiword") is not None
if not ANTIWORD_AVAILABLE:
    logger.debug("antiword command not found. Install it for better .doc extraction.")

# Check for LibreOffice command-line tool
LIBREOFFICE_AVAILABLE = shutil.which("soffice") is not None or shutil.which("libreoffice") is not None
if not LIBREOFFICE_AVAILABLE:
    logger.debug("LibreOffice not found. Install it for most reliable .doc conversion.")


class ResumeParser:
    """Service for parsing resume files and extracting text."""
    
    def __init__(self):
        """Initialize ResumeParser."""
        pass
    
    def _convert_image_to_pdf(self, file_content: bytes, filename: str) -> bytes:
        """
        Convert image file to PDF for better OCR extraction.
        
        Args:
            file_content: The binary content of the image file
            filename: Name of the file
            
        Returns:
            PDF content as bytes
        """
        if not OCR_AVAILABLE or not PYMUPDF_AVAILABLE:
            raise ValueError("PyMuPDF required for image to PDF conversion")
        
        try:
            # Load image
            image = Image.open(BytesIO(file_content))
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Convert to PDF using PyMuPDF
            img_bytes = BytesIO()
            image.save(img_bytes, format='PDF')
            img_bytes.seek(0)
            
            # Use PyMuPDF to create a proper PDF
            pdf_doc = fitz.open(stream=img_bytes.read(), filetype="pdf")
            pdf_bytes = pdf_doc.tobytes()
            pdf_doc.close()
            
            logger.info(
                f"✅ Converted image {filename} to PDF for OCR extraction",
                extra={"file_name": filename}
            )
            return pdf_bytes
        except Exception as e:
            logger.error(f"Failed to convert image to PDF: {e}", extra={"file_name": filename})
            raise ValueError(f"Failed to convert image to PDF: {e}")
    
    def _convert_html_to_pdf(self, file_content: bytes, filename: str) -> bytes:
        """
        Convert HTML file to PDF for better OCR extraction.
        Uses PyMuPDF to render HTML to PDF.
        
        Args:
            file_content: The binary content of the HTML file
            filename: Name of the file
            
        Returns:
            PDF content as bytes
        """
        if not PYMUPDF_AVAILABLE:
            raise ValueError("PyMuPDF required for HTML to PDF conversion")
        
        try:
            # Decode HTML content
            html_content = file_content.decode('utf-8', errors='ignore')
            
            # Create PDF from HTML using PyMuPDF
            pdf_doc = fitz.open()  # Create new PDF
            page = pdf_doc.new_page()
            
            # Insert HTML content
            page.insert_html(html_content)
            
            # Convert to bytes
            pdf_bytes = pdf_doc.tobytes()
            pdf_doc.close()
            
            logger.info(
                f"✅ Converted HTML {filename} to PDF for OCR extraction",
                extra={"file_name": filename}
            )
            return pdf_bytes
        except Exception as e:
            logger.warning(
                f"PyMuPDF HTML to PDF conversion failed, trying alternative method: {e}",
                extra={"file_name": filename}
            )
            # Alternative: Convert HTML to image then to PDF
            try:
                if OCR_AVAILABLE and HTML_PARSING_AVAILABLE:
                    # Parse HTML and extract text, then create PDF from text
                    soup = BeautifulSoup(html_content, 'html.parser')
                    text = soup.get_text(separator='\n', strip=True)
                    
                    # Create PDF with text
                    pdf_doc = fitz.open()
                    page = pdf_doc.new_page()
                    page.insert_text((50, 50), text, fontsize=11)
                    pdf_bytes = pdf_doc.tobytes()
                    pdf_doc.close()
                    
                    logger.info(
                        f"✅ Converted HTML {filename} to PDF using text extraction",
                        extra={"file_name": filename}
                    )
                    return pdf_bytes
            except Exception as alt_error:
                logger.error(f"Alternative HTML to PDF conversion also failed: {alt_error}")
                raise ValueError(f"Failed to convert HTML to PDF: {e}")
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from uploaded file based on extension.
        Supports PDF, DOCX, DOC, TXT, images (JPG, PNG), and HTML files.
        
        Args:
            file_content: The binary content of the file
            filename: Name of the file (used to determine file type)
        
        Returns:
            Extracted text content as string
        
        Raises:
            ValueError: If file type is not supported or extraction fails
        """
        try:
            filename_lower = filename.lower()
            
            # PDF files - try regular extraction first, fallback to OCR if needed
            if filename_lower.endswith('.pdf'):
                return self._extract_pdf_text(file_content, filename)
            
            # DOCX files
            elif filename_lower.endswith('.docx'):
                return self._extract_docx_text(file_content, filename)
            
            # DOC files
            elif filename_lower.endswith('.doc'):
                return self._extract_doc_text(file_content)
            
            # Image files - try converting to PDF first for better OCR, then extract
            elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif')):
                # First try direct OCR extraction
                try:
                    image_text = self._extract_image_text(file_content, filename)
                    if image_text and len(image_text.strip()) >= 50:
                        return image_text
                except Exception as img_error:
                    logger.debug(f"Direct image OCR failed, trying PDF conversion: {img_error}")
                
                # If direct OCR failed or returned insufficient text, convert to PDF and extract
                if PYMUPDF_AVAILABLE:
                    try:
                        logger.info(
                            f"🔄 Converting image {filename} to PDF for better OCR extraction",
                            extra={"file_name": filename}
                        )
                        pdf_content = self._convert_image_to_pdf(file_content, filename)
                        # Extract from PDF (which will use OCR if needed)
                        return self._extract_pdf_text(pdf_content, filename.replace('.jpg', '.pdf').replace('.jpeg', '.pdf').replace('.png', '.pdf'))
                    except Exception as pdf_conv_error:
                        logger.warning(f"Image to PDF conversion failed, using direct OCR: {pdf_conv_error}")
                        # Fallback to direct OCR
                        return self._extract_image_text(file_content, filename)
                else:
                    # No PDF conversion available, use direct OCR
                    return self._extract_image_text(file_content, filename)
            
            # HTML files - try converting to PDF first for better extraction, then extract
            elif filename_lower.endswith(('.html', '.htm')):
                # First try direct HTML extraction
                try:
                    html_text = self._extract_html_text(file_content, filename)
                    if html_text and len(html_text.strip()) >= 50:
                        return html_text
                except Exception as html_error:
                    logger.debug(f"Direct HTML extraction failed, trying PDF conversion: {html_error}")
                
                # If direct HTML extraction failed or returned insufficient text, convert to PDF and extract
                if PYMUPDF_AVAILABLE:
                    try:
                        logger.info(
                            f"🔄 Converting HTML {filename} to PDF for better OCR extraction",
                            extra={"file_name": filename}
                        )
                        pdf_content = self._convert_html_to_pdf(file_content, filename)
                        # Extract from PDF (which will use OCR if needed)
                        return self._extract_pdf_text(pdf_content, filename.replace('.html', '.pdf').replace('.htm', '.pdf'))
                    except Exception as pdf_conv_error:
                        logger.warning(f"HTML to PDF conversion failed, using direct HTML extraction: {pdf_conv_error}")
                        # Fallback to direct HTML extraction
                        return self._extract_html_text(file_content, filename)
                else:
                    # No PDF conversion available, use direct HTML extraction
                    return self._extract_html_text(file_content, filename)
            
            # Text files
            elif filename_lower.endswith('.txt'):
                return file_content.decode('utf-8', errors='ignore')
            
            else:
                # Try as text for unknown extensions
                return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            # Safe logging - avoid using reserved LogRecord attributes
            error_msg = f"Error extracting text from {filename}: {e}"
            print(f"[ERROR] {error_msg}")  # Print to console for visibility
            
            # Try fallback extraction for HTML and image files
            filename_lower = filename.lower()
            if filename_lower.endswith(('.html', '.htm')):
                try:
                    logger.info(f"🔄 Attempting fallback HTML extraction for {filename}")
                    # Try simple text extraction
                    text = file_content.decode('utf-8', errors='ignore')
                    
                    # Remove forwarding headers (HTML-specific filtering)
                    forwarding_patterns = [
                        r'(?i)Forwarded\s+By:.*?\n',
                        r'(?i)To:\s*\[.*?\]\s*\n',
                        r'(?i)From:.*?\n',
                        r'(?i)Resume\s+Link:.*?\n',
                        r'(?i)Comments:.*?\n',
                        r'(?i)This\s+resume\s+has\s+been\s+forwarded.*?\n',
                        r'(?i)This\s+email\s+was\s+sent.*?\n',
                    ]
                    for pattern in forwarding_patterns:
                        text = re.sub(pattern, '', text, flags=re.DOTALL)
                    
                    # Remove HTML tags with regex
                    text = re.sub(r'<[^>]+>', ' ', text)
                    # Decode HTML entities
                    try:
                        text = html.unescape(text)
                    except:
                        pass
                    text = ' '.join(text.split())
                    if text and len(text.strip()) > 20:
                        normalized_text = normalize_text(text) or text
                        logger.info(f"✅ Fallback HTML extraction succeeded for {filename}")
                        return normalized_text
                except Exception as fallback_error:
                    logger.debug(f"Fallback HTML extraction failed: {fallback_error}")
            
            try:
                # Use safe_extra to prevent LogRecord conflicts
                safe_extras = safe_extra({"error": str(e), "file_name": filename})
                logger.error(error_msg, extra=safe_extras)
            except Exception as log_error:
                # If logging fails, at least print the error
                print(f"[CRITICAL] Logging failed: {log_error}")
                print(f"[CRITICAL] Original error: {error_msg}")
            raise ValueError(f"Failed to extract text from file: {e}")
    
    def _extract_pdf_text(self, file_content: bytes, filename: str = "resume.pdf") -> str:
        """
        Extract text from PDF file.
        First tries regular text extraction, ALWAYS tries OCR for image-based PDFs.
        For image-based PDFs, OCR is the primary method.
        """
        normalized_text = ""
        text_length = 0
        word_count = 0
        
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as page_error:
                    logger.warning(
                        f"Failed to extract text from a page in {filename}: {page_error}",
                        extra={"file_name": filename, "page_error": str(page_error)}
                    )
                    continue
            
            raw_text = "\n".join(text_parts)
            # Normalize whitespace (remove extra spaces, normalize line breaks)
            normalized_text = normalize_text(raw_text) or raw_text
            
            # Check if extracted text is minimal (likely a scanned PDF)
            # For image-based PDFs, ALWAYS try OCR - it's the primary extraction method
            text_length = len(normalized_text.strip())
            word_count = len(re.findall(r'\b\w+\b', normalized_text)) if normalized_text else 0
            
            # Detect if PDF is image-based (scanned PDF):
            # 1. Very little text extracted (< 100 chars)
            # 2. Very few words (< 10 words)
            # 3. Text is mostly whitespace or special characters
            is_likely_image_based = (
                text_length < 100 or 
                word_count < 10 or
                (text_length > 0 and word_count == 0)  # Has characters but no words
            )
            
            # ALWAYS try OCR for PDFs if available (OCR is better for image-based PDFs)
            # For image-based PDFs, OCR is the primary extraction method
            # Always attempt OCR if text is minimal OR if we suspect it's image-based
            if is_likely_image_based:  # More aggressive: try OCR for any suspicious PDF
                # Try OCR first if available
                if OCR_AVAILABLE and (PDF2IMAGE_AVAILABLE or PYMUPDF_AVAILABLE):
                    logger.info(
                        f"📄 PDF text extraction: {text_length} chars, {word_count} words. "
                        f"Detected as image-based PDF. Attempting OCR: {filename}",
                        extra={
                            "file_name": filename, 
                            "text_length": text_length,
                            "word_count": word_count,
                            "is_image_based": is_likely_image_based
                        }
                    )
                    
                    try:
                        ocr_text = self._extract_scanned_pdf_text(file_content, filename)
                        if ocr_text and len(ocr_text.strip()) > 0:
                            # OCR succeeded - use OCR text (it's better for image-based PDFs)
                            ocr_length = len(ocr_text.strip())
                            ocr_word_count = len(re.findall(r'\b\w+\b', ocr_text)) if ocr_text else 0
                            
                            # ALWAYS prefer OCR text if:
                            # 1. Regular extraction was minimal (< 100 chars OR < 10 words)
                            # 2. OCR got more text
                            # 3. OCR got more words (even if same length)
                            if (is_likely_image_based or 
                                ocr_length > text_length or 
                                (ocr_length >= text_length and ocr_word_count > word_count)):
                                logger.info(
                                    f"✅ OCR extraction SUCCESS for {filename}: "
                                    f"extracted {ocr_length} chars, {ocr_word_count} words "
                                    f"(vs {text_length} chars, {word_count} words from regular extraction)",
                                    extra={
                                        "file_name": filename, 
                                        "ocr_text_length": ocr_length,
                                        "ocr_word_count": ocr_word_count,
                                        "regular_text_length": text_length,
                                        "regular_word_count": word_count
                                    }
                                )
                                return ocr_text
                            else:
                                # OCR got less text but regular extraction was good, use regular
                                logger.info(
                                    f"Using regular extraction for {filename} "
                                    f"(OCR: {ocr_length} chars, Regular: {text_length} chars)",
                                    extra={"file_name": filename, "ocr_text_length": ocr_length, "regular_text_length": text_length}
                                )
                        else:
                            logger.warning(
                                f"⚠️ OCR extraction returned empty text for PDF: {filename}",
                                extra={"file_name": filename}
                            )
                            # If OCR returned empty and regular extraction was also minimal, return empty
                            if is_likely_image_based:
                                logger.warning(
                                    f"Both OCR and regular extraction failed for image-based PDF: {filename}",
                                    extra={"file_name": filename, "regular_text_length": text_length, "regular_word_count": word_count}
                                )
                                return ""
                    except Exception as ocr_error:
                        logger.error(
                            f"❌ OCR extraction failed for PDF {filename}: {ocr_error}. "
                            f"Please ensure Tesseract OCR and poppler are installed.",
                            extra={"file_name": filename, "error": str(ocr_error)},
                            exc_info=True
                        )
                        # If OCR fails and regular extraction was minimal, try PyMuPDF text extraction
                        if PYMUPDF_AVAILABLE:
                            logger.info(
                                f"OCR failed, trying PyMuPDF built-in text extraction for {filename}",
                                extra={"file_name": filename}
                            )
                            try:
                                pymupdf_doc = fitz.open(stream=file_content, filetype="pdf")
                                pymupdf_text_parts = []
                                for page_num in range(len(pymupdf_doc)):
                                    page = pymupdf_doc[page_num]
                                    page_text = page.get_text()
                                    if page_text:
                                        pymupdf_text_parts.append(page_text)
                                pymupdf_doc.close()
                                
                                if pymupdf_text_parts:
                                    pymupdf_text = "\n".join(pymupdf_text_parts)
                                    pymupdf_length = len(pymupdf_text.strip())
                                    if pymupdf_length > text_length:
                                        logger.info(
                                            f"✅ PyMuPDF text extraction SUCCESS for {filename}: "
                                            f"extracted {pymupdf_length} chars",
                                            extra={"file_name": filename, "pymupdf_text_length": pymupdf_length}
                                        )
                                        return normalize_text(pymupdf_text) or pymupdf_text
                            except Exception as pymupdf_error:
                                logger.warning(
                                    f"PyMuPDF text extraction also failed: {pymupdf_error}",
                                    extra={"file_name": filename, "error": str(pymupdf_error)}
                                )
                elif not OCR_AVAILABLE and PYMUPDF_AVAILABLE:
                    # OCR not available, but try PyMuPDF text extraction
                    logger.info(
                        f"OCR not available, trying PyMuPDF built-in text extraction for {filename}",
                        extra={"file_name": filename}
                    )
                    try:
                        pdf_file.seek(0)  # Reset file pointer
                        pymupdf_doc = fitz.open(stream=file_content, filetype="pdf")
                        pymupdf_text_parts = []
                        for page_num in range(len(pymupdf_doc)):
                            page = pymupdf_doc[page_num]
                            page_text = page.get_text()
                            if page_text:
                                pymupdf_text_parts.append(page_text)
                        pymupdf_doc.close()
                        
                        if pymupdf_text_parts:
                            pymupdf_text = "\n".join(pymupdf_text_parts)
                            pymupdf_length = len(pymupdf_text.strip())
                            if pymupdf_length > text_length:
                                logger.info(
                                    f"✅ PyMuPDF text extraction SUCCESS for {filename}: "
                                    f"extracted {pymupdf_length} chars (vs {text_length} from regular extraction)",
                                    extra={"file_name": filename, "pymupdf_text_length": pymupdf_length, "regular_text_length": text_length}
                                )
                                return normalize_text(pymupdf_text) or pymupdf_text
                    except Exception as pymupdf_error:
                        logger.warning(
                            f"PyMuPDF text extraction failed: {pymupdf_error}",
                            extra={"file_name": filename, "error": str(pymupdf_error)}
                        )
            
            # Return the best text we have (either OCR, PyMuPDF, or regular extraction)
            if normalized_text and text_length >= 50:
                return normalized_text
            elif normalized_text:
                logger.warning(
                    f"PyPDF2 extracted only {text_length} chars from {filename}, trying PyMuPDF fallback",
                    extra={"file_name": filename, "text_length": text_length}
                )
        except Exception as e:
            logger.warning(
                f"PyPDF2 extraction failed for {filename}: {e}, trying PyMuPDF fallback",
                extra={"file_name": filename, "error": str(e)}
            )
        
        # If PyPDF2 failed or returned insufficient text, try PyMuPDF as fallback
        if PYMUPDF_AVAILABLE:
            try:
                logger.info(
                    f"Attempting PyMuPDF text extraction for {filename}",
                    extra={"file_name": filename}
                )
                pymupdf_doc = fitz.open(stream=file_content, filetype="pdf")
                pymupdf_text_parts = []
                for page_num in range(len(pymupdf_doc)):
                    try:
                        page = pymupdf_doc[page_num]
                        page_text = page.get_text()
                        if page_text:
                            pymupdf_text_parts.append(page_text)
                    except Exception as page_error:
                        logger.warning(
                            f"Failed to extract text from page {page_num + 1} with PyMuPDF: {page_error}",
                            extra={"file_name": filename, "page_num": page_num + 1}
                        )
                        continue
                pymupdf_doc.close()
                
                if pymupdf_text_parts:
                    pymupdf_text = "\n".join(pymupdf_text_parts)
                    pymupdf_normalized = normalize_text(pymupdf_text) or pymupdf_text
                    pymupdf_length = len(pymupdf_normalized.strip())
                    pymupdf_word_count = len(re.findall(r'\b\w+\b', pymupdf_normalized)) if pymupdf_normalized else 0
                    
                    if pymupdf_length > text_length or (pymupdf_length >= 50 and text_length < 50):
                        logger.info(
                            f"✅ PyMuPDF text extraction SUCCESS for {filename}: "
                            f"extracted {pymupdf_length} chars, {pymupdf_word_count} words",
                            extra={
                                "file_name": filename,
                                "pymupdf_text_length": pymupdf_length,
                                "pymupdf_word_count": pymupdf_word_count,
                                "regular_text_length": text_length
                            }
                        )
                        return pymupdf_normalized
                    else:
                        logger.debug(
                            f"PyMuPDF extracted {pymupdf_length} chars (not better than PyPDF2's {text_length})",
                            extra={"file_name": filename, "pymupdf_length": pymupdf_length, "text_length": text_length}
                        )
            except Exception as pymupdf_error:
                logger.warning(
                    f"PyMuPDF text extraction also failed for {filename}: {pymupdf_error}",
                    extra={"file_name": filename, "error": str(pymupdf_error)}
                )
        
        # If we have any text from PyPDF2, return it (even if minimal)
        if normalized_text:
            return normalized_text
        
        # If all methods failed, raise error
        logger.error(
            f"❌ All PDF text extraction methods failed for {filename}",
            extra={"file_name": filename}
        )
        raise ValueError(f"Failed to extract text from PDF: All extraction methods failed")
    
    def _extract_docx_text(self, file_content: bytes, filename: str = "resume.docx") -> str:
        """
        Extract text from DOCX file.
        Also extracts text from embedded images using OCR if available.
        Enhanced to handle more edge cases and better text extraction.
        """
        try:
            doc_file = BytesIO(file_content)
            doc = Document(doc_file)
            text_parts = []
            
            # Extract from paragraphs (including runs with different formatting)
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text_parts.append(para_text)
                # Also check individual runs (sometimes text is split across runs)
                for run in paragraph.runs:
                    run_text = run.text.strip()
                    if run_text and run_text not in para_text:
                        text_parts.append(run_text)
            
            # Extract from tables (contact info is often in tables)
            # Tables are critical for resumes - contact info is often in header tables
            for table_idx, table in enumerate(doc.tables):
                table_text_parts = []
                for row in table.rows:
                    row_text_parts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text_parts.append(cell_text)
                            # Also check cell paragraphs individually
                            for para in cell.paragraphs:
                                para_text = para.text.strip()
                                if para_text and para_text not in cell_text:
                                    row_text_parts.append(para_text)
                    if row_text_parts:
                        # Join row cells with separator for better readability
                        table_text_parts.append(' | '.join(row_text_parts))
                if table_text_parts:
                    text_parts.extend(table_text_parts)
                    logger.debug(f"Extracted {len(table_text_parts)} rows from table {table_idx+1} in {filename}")
            
            # Extract from headers and footers (contact info is often in headers)
            # Headers are VERY important for contact information
            for section_idx, section in enumerate(doc.sections):
                # Header
                if section.header:
                    header_parts = []
                    for paragraph in section.header.paragraphs:
                        header_text = paragraph.text.strip()
                        if header_text:
                            header_parts.append(header_text)
                        # Also check runs in header
                        for run in paragraph.runs:
                            run_text = run.text.strip()
                            if run_text and run_text not in header_text:
                                header_parts.append(run_text)
                    if header_parts:
                        text_parts.extend(header_parts)
                        logger.debug(f"Extracted {len(header_parts)} header parts from section {section_idx+1} in {filename}")
                
                # Footer
                if section.footer:
                    footer_parts = []
                    for paragraph in section.footer.paragraphs:
                        footer_text = paragraph.text.strip()
                        if footer_text:
                            footer_parts.append(footer_text)
                        # Also check runs in footer
                        for run in paragraph.runs:
                            run_text = run.text.strip()
                            if run_text and run_text not in footer_text:
                                footer_parts.append(run_text)
                    if footer_parts:
                        text_parts.extend(footer_parts)
                        logger.debug(f"Extracted {len(footer_parts)} footer parts from section {section_idx+1} in {filename}")
            
            # Extract text from embedded images using OCR (if available)
            # This is important for resumes where contact info might be in images
            if OCR_AVAILABLE:
                try:
                    # DOCX files are ZIP archives containing XML and media files
                    # Create a new BytesIO from the original content since doc_file position may have changed
                    doc_file_for_zip = BytesIO(file_content)
                    docx_zip = zipfile.ZipFile(doc_file_for_zip)
                    
                    # Look for images in the media folder
                    image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/') and 
                                  f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif'))]
                    
                    if image_files:
                        logger.info(
                            f"Found {len(image_files)} embedded image(s) in DOCX, attempting OCR",
                            extra={"image_count": len(image_files)}
                        )
                        
                        for image_path in image_files:
                            try:
                                # Extract image from DOCX
                                image_data = docx_zip.read(image_path)
                                
                                # Load image and perform OCR
                                image = Image.open(BytesIO(image_data))
                                
                                # Convert to RGB if necessary
                                if image.mode != 'RGB':
                                    image = image.convert('RGB')
                                
                                # Pre-process image for better OCR
                                img_array = np.array(image)
                                img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                                gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                                thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                                processed_image = Image.fromarray(thresh)
                                
                                # Perform OCR
                                ocr_text = pytesseract.image_to_string(processed_image, lang='eng')
                                if ocr_text and ocr_text.strip():
                                    text_parts.append(ocr_text.strip())
                                    logger.debug(
                                        f"Extracted {len(ocr_text.strip())} characters from embedded image: {image_path}",
                                        extra={"image_path": image_path, "ocr_text_length": len(ocr_text.strip())}
                                    )
                            except Exception as img_error:
                                logger.debug(f"Failed to extract text from embedded image {image_path}: {img_error}")
                                continue
                except Exception as ocr_error:
                    logger.debug(f"Failed to extract images from DOCX for OCR: {ocr_error}")
            
            raw_text = "\n".join(text_parts)
            
            # Check if we got any text
            if not raw_text or len(raw_text.strip()) < 10:
                logger.warning(
                    f"⚠️ DOCX extraction returned minimal text for {filename} ({len(raw_text.strip())} chars). "
                    f"Trying alternative extraction methods...",
                    extra={"file_name": filename, "text_length": len(raw_text.strip())}
                )
                
                # Try extracting from XML directly as fallback
                try:
                    doc_file_for_xml = BytesIO(file_content)
                    docx_zip = zipfile.ZipFile(doc_file_for_xml)
                    
                    # Read main document XML
                    if 'word/document.xml' in docx_zip.namelist():
                        import xml.etree.ElementTree as ET
                        xml_content = docx_zip.read('word/document.xml')
                        root = ET.fromstring(xml_content)
                        
                        # Extract text from all text nodes
                        xml_text_parts = []
                        for elem in root.iter():
                            if elem.text and elem.text.strip():
                                xml_text_parts.append(elem.text.strip())
                        
                        if xml_text_parts:
                            xml_text = ' '.join(xml_text_parts)
                            if len(xml_text.strip()) > len(raw_text.strip()):
                                logger.info(f"✅ XML extraction found more text ({len(xml_text.strip())} chars) for {filename}")
                                raw_text = xml_text
                except Exception as xml_error:
                    logger.debug(f"XML extraction fallback failed: {xml_error}")
            
            # Normalize whitespace (remove extra spaces, normalize line breaks)
            normalized_text = normalize_text(raw_text) or raw_text
            
            if normalized_text and len(normalized_text.strip()) > 10:
                logger.info(
                    f"✅ DOCX text extraction completed for {filename} (extracted {len(normalized_text.strip())} chars)",
                    extra={"file_name": filename, "text_length": len(normalized_text.strip())}
                )
            else:
                logger.warning(
                    f"⚠️ DOCX extraction returned insufficient text for {filename} ({len(normalized_text.strip()) if normalized_text else 0} chars)",
                    extra={"file_name": filename, "text_length": len(normalized_text.strip()) if normalized_text else 0}
                )
            
            return normalized_text
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {filename}: {e}", extra={"error": str(e), "file_name": filename})
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    def _extract_doc_text(self, file_content: bytes) -> str:
        """
        Extract text from DOC file (older Microsoft Word format).
        Uses methods in order of reliability:
        1. Apache Tika (PRIMARY - Currently Working)
        2. LibreOffice headless conversion (if available)
        3. antiword (if available)
        4. python-docx fallback (might work for some files)
        5. olefile (basic binary extraction - fallback)
        """
        # Create temporary file for .doc content
        with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
            temp_file.write(file_content)
            temp_doc_path = temp_file.name
        
        try:
            # Method 1: Apache Tika (PRIMARY METHOD - Currently Working)
            if TIKA_AVAILABLE:
                try:
                    tika_msg = (
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$  USING APACHE TIKA  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$  Extracting text from .doc file using Apache Tika\n"
                        "$$$  This is the PRIMARY method for processing .doc files\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                    )
                    print(tika_msg)
                    logger.info(tika_msg)
                    parsed = tika_parser.from_file(temp_doc_path)
                    if parsed and 'content' in parsed and parsed['content']:
                        text = parsed['content'].strip()
                        if text:
                            # Normalize whitespace (remove extra spaces, normalize line breaks)
                            normalized_text = normalize_text(text) or text
                            success_msg = (
                                "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                "$$$$$$$$$$$$$$$$$$$$$$  APACHE TIKA SUCCESS  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                f"$$$  Successfully extracted {len(normalized_text)} characters using Apache Tika\n"
                                "$$$  METHOD USED: Apache Tika (tika-python library)\n"
                                "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                            )
                            print(success_msg)
                            logger.info(
                                success_msg,
                                extra={"extraction_method": "apache_tika", "text_length": len(normalized_text)}
                            )
                            return normalized_text
                except Exception as tika_error:
                    error_msg = f"Apache Tika extraction failed: {tika_error}"
                    print(f"[WARNING] {error_msg}")
                    logger.warning(error_msg)
            
            # Method 2: LibreOffice headless conversion (if available)
            if LIBREOFFICE_AVAILABLE:
                try:
                    lo_msg = (
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$$$$$$$$$$$$$$$$$$$$  USING LIBREOFFICE  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$  Converting .doc to .docx using LibreOffice headless\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                    )
                    print(lo_msg)
                    logger.info(lo_msg)
                    # Create temp directory for output
                    with tempfile.TemporaryDirectory() as temp_dir:
                        # Use LibreOffice to convert .doc to .docx
                        libreoffice_cmd = shutil.which("soffice") or shutil.which("libreoffice")
                        cmd = [
                            libreoffice_cmd,
                            "--headless",
                            "--convert-to", "docx",
                            "--outdir", temp_dir,
                            temp_doc_path
                        ]
                        
                        result = subprocess.run(
                            cmd,
                            capture_output=True,
                            text=True,
                            timeout=30
                        )
                        
                        if result.returncode == 0:
                            # Find the converted .docx file
                            converted_docx = os.path.join(temp_dir, os.path.basename(temp_doc_path).replace('.doc', '.docx'))
                            if os.path.exists(converted_docx):
                                # Read the converted .docx and extract text
                                with open(converted_docx, 'rb') as f:
                                    docx_content = f.read()
                                text = self._extract_docx_text(docx_content)
                                if text.strip():
                                    success_msg = (
                                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                        "$$$$$$$$$$$$$$$$$$$$$$  LIBREOFFICE SUCCESS  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                        f"$$$  Successfully extracted {len(text)} characters using LibreOffice\n"
                                        "$$$  METHOD USED: LibreOffice (converted .doc to .docx, then extracted)\n"
                                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                                    )
                                    print(success_msg)
                                    logger.info(
                                        success_msg,
                                        extra={"extraction_method": "libreoffice", "text_length": len(text)}
                                    )
                                    return text
                except subprocess.TimeoutExpired:
                    logger.warning("LibreOffice conversion timed out")
                except Exception as lo_error:
                    logger.debug(f"LibreOffice conversion failed: {lo_error}")
            
            # Method 3: antiword (if available)
            if ANTIWORD_AVAILABLE:
                try:
                    antiword_msg = (
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$  USING ANTIWORD  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$  Extracting text from .doc file using antiword\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                    )
                    print(antiword_msg)
                    logger.info(antiword_msg)
                    result = subprocess.run(
                        ["antiword", temp_doc_path],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        # Normalize whitespace (remove extra spaces, normalize line breaks)
                        normalized_text = normalize_text(result.stdout) or result.stdout
                        success_msg = (
                            "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                            "$$$$$$$$$$$$$$$$$$$$$$  ANTIWORD SUCCESS  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                            f"$$$  Successfully extracted {len(normalized_text)} characters using antiword\n"
                            "$$$  METHOD USED: antiword (command-line tool)\n"
                            "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                        )
                        print(success_msg)
                        logger.info(
                            success_msg,
                            extra={"extraction_method": "antiword", "text_length": len(normalized_text)}
                        )
                        return normalized_text
                except subprocess.TimeoutExpired:
                    logger.warning("antiword extraction timed out")
                except Exception as aw_error:
                    logger.debug(f"antiword extraction failed: {aw_error}")
            
            # Method 4: Try python-docx as fallback (might work for some .doc files that are actually .docx)
            try:
                logger.debug("Attempting .doc extraction using python-docx fallback")
                doc_file = BytesIO(file_content)
                doc = Document(doc_file)
                text_parts = []
                for paragraph in doc.paragraphs:
                    text_parts.append(paragraph.text)
                # Also try to extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                extracted_text = "\n".join(text_parts)
                if extracted_text.strip():
                    # Normalize whitespace (remove extra spaces, normalize line breaks)
                    normalized_text = normalize_text(extracted_text) or extracted_text
                    logger.info("Successfully extracted .doc file using python-docx fallback")
                    return normalized_text
            except Exception as fallback_error:
                logger.debug(f"python-docx fallback failed: {fallback_error}")
            
            # Method 5: Try to extract using olefile (for binary .doc files - last resort)
            if OLEFILE_AVAILABLE:
                try:
                    logger.debug("Attempting .doc extraction using olefile (last resort)")
                    # .doc files are OLE compound documents
                    ole = olefile.OleFileIO(BytesIO(file_content))
                    # Try to find WordDocument stream
                    if ole.exists('WordDocument'):
                        stream = ole.openstream('WordDocument')
                        # Read and try to extract text (basic extraction)
                        data = stream.read()
                        # Simple text extraction from binary data
                        # This is a basic approach - look for readable text
                        text_chunks = []
                        current_chunk = b""
                        for byte in data:
                            if 32 <= byte <= 126 or byte in [9, 10, 13]:  # Printable ASCII
                                current_chunk += bytes([byte])
                            else:
                                if len(current_chunk) > 3:
                                    try:
                                        text_chunks.append(current_chunk.decode('ascii', errors='ignore'))
                                    except:
                                        pass
                                current_chunk = b""
                        if current_chunk:
                            try:
                                text_chunks.append(current_chunk.decode('ascii', errors='ignore'))
                            except:
                                pass
                        ole.close()
                        extracted_text = "\n".join(text_chunks)
                        if extracted_text.strip():
                            # Normalize whitespace (remove extra spaces, normalize line breaks)
                            normalized_text = normalize_text(extracted_text) or extracted_text
                            logger.info("Successfully extracted .doc file using olefile")
                            return normalized_text
                    ole.close()
                except Exception as ole_error:
                    logger.debug(f"olefile extraction failed: {ole_error}")
            
            # Last resort: Try to extract as plain text (basic binary extraction)
            try:
                logger.debug("Attempting basic text extraction from .doc file as last resort")
                # Try to extract readable text from binary data
                text_chunks = []
                # Look for readable ASCII text sequences
                current_chunk = b""
                for byte in file_content:
                    if 32 <= byte <= 126 or byte in [9, 10, 13]:  # Printable ASCII
                        current_chunk += bytes([byte])
                    else:
                        if len(current_chunk) > 10:  # Only keep meaningful chunks
                            try:
                                decoded = current_chunk.decode('ascii', errors='ignore')
                                if decoded.strip():
                                    text_chunks.append(decoded.strip())
                            except:
                                pass
                        current_chunk = b""
                
                # Add last chunk
                if len(current_chunk) > 10:
                    try:
                        decoded = current_chunk.decode('ascii', errors='ignore')
                        if decoded.strip():
                            text_chunks.append(decoded.strip())
                    except:
                        pass
                
                if text_chunks:
                    extracted_text = "\n".join(text_chunks)
                    normalized_text = normalize_text(extracted_text) or extracted_text
                    if normalized_text and len(normalized_text.strip()) > 20:
                        logger.info(
                            f"✅ Successfully extracted .doc file using basic binary extraction (extracted {len(normalized_text.strip())} chars)",
                            extra={"extraction_method": "binary_fallback", "text_length": len(normalized_text.strip())}
                        )
                        return normalized_text
            except Exception as binary_error:
                logger.debug(f"Basic binary extraction also failed: {binary_error}")
            
            # If all methods fail, raise an error with helpful message
            raise ValueError(
                "Failed to extract text from .doc file using all available methods.\n"
                "Installation options:\n"
                "1. Apache Tika: pip install tika (REQUIRES Java runtime) - PRIMARY METHOD\n"
                "2. LibreOffice (headless): Most reliable for production\n"
                "   - Windows: Download from https://www.libreoffice.org/\n"
                "   - Linux: sudo apt-get install libreoffice\n"
                "3. antiword: Good for plain-text extraction\n"
                "   - Windows: Download from http://www.winfield.demon.nl/\n"
                "   - Linux: sudo apt-get install antiword\n"
                "4. olefile: pip install olefile (basic support, already installed)\n"
                "5. Convert .doc files to .docx format before processing"
            )
        finally:
            # Clean up temp file
            if os.path.exists(temp_doc_path):
                os.unlink(temp_doc_path)
    
    def _extract_image_text(self, file_content: bytes, filename: str = "resume.jpg") -> str:
        """
        Extract text from image file using OCR (Tesseract) with enhanced preprocessing.
        Includes deskew, noise removal, and DPI increase for better accuracy.
        
        Args:
            file_content: The binary content of the image file
            filename: Name of the file (for logging)
        
        Returns:
            Extracted text content as string
        
        Raises:
            ValueError: If OCR is not available or extraction fails
        """
        if not OCR_AVAILABLE:
            raise ValueError(
                "OCR libraries not available. Install required packages:\n"
                "pip install pytesseract pillow opencv-python\n"
                "Also install Tesseract OCR:\n"
                "  - Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki\n"
                "  - Linux: sudo apt-get install tesseract-ocr\n"
                "  - macOS: brew install tesseract"
            )
        
        try:
            # Load image from bytes
            image = Image.open(BytesIO(file_content))
            original_size = image.size
            
            # Convert to RGB if necessary (some images are RGBA or other formats)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Increase DPI/Resolution for better OCR (300 DPI recommended)
            # Scale up image if it's too small
            if image.size[0] < 1200 or image.size[1] < 1200:
                scale_factor = max(1200 / image.size[0], 1200 / image.size[1])
                new_size = (int(image.size[0] * scale_factor), int(image.size[1] * scale_factor))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
                logger.debug(f"Scaled image from {original_size} to {new_size} for better OCR")
            
            # Pre-process image for better OCR accuracy
            # Convert PIL image to OpenCV format (numpy array)
            img_array = np.array(image)
            img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            
            # Convert to grayscale
            gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
            
            # Noise removal using median blur (removes salt and pepper noise)
            denoised = cv2.medianBlur(gray, 3)
            
            # Deskew image (correct rotation)
            try:
                # Find angle using Hough transform
                coords = np.column_stack(np.where(denoised > 0))
                if len(coords) > 0:
                    angle = cv2.minAreaRect(coords)[-1]
                    if angle < -45:
                        angle = -(90 + angle)
                    else:
                        angle = -angle
                    # Only apply if angle is significant (> 0.5 degrees)
                    if abs(angle) > 0.5:
                        (h, w) = denoised.shape[:2]
                        center = (w // 2, h // 2)
                        M = cv2.getRotationMatrix2D(center, angle, 1.0)
                        denoised = cv2.warpAffine(denoised, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
                        logger.debug(f"Deskewed image by {angle:.2f} degrees")
            except Exception as deskew_error:
                logger.debug(f"Deskew failed (non-critical): {deskew_error}")
                # Continue without deskew
            
            # Apply thresholding to get binary image (improves OCR accuracy)
            # Try multiple thresholding methods and use the best result
            methods = [
                ("OTSU", cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]),
                ("ADAPTIVE", cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)),
                ("SIMPLE", cv2.threshold(denoised, 150, 255, cv2.THRESH_BINARY)[1]),
            ]
            
            best_text = ""
            best_length = 0
            
            for method_name, thresh in methods:
                try:
                    # Convert back to PIL Image for pytesseract
                    processed_image = Image.fromarray(thresh)
                    
                    # Perform OCR with different configurations
                    configs = [
                        '--psm 6',  # Assume uniform block of text
                        '--psm 11',  # Sparse text
                        '--psm 12',  # Sparse text with OSD
                    ]
                    
                    for config in configs:
                        try:
                            text = pytesseract.image_to_string(processed_image, lang='eng', config=config)
                            if len(text.strip()) > best_length:
                                best_text = text
                                best_length = len(text.strip())
                        except:
                            continue
                except Exception as e:
                    logger.debug(f"OCR method {method_name} failed: {e}")
                    continue
            
            # If no text found with advanced methods, try basic method
            if not best_text or best_length < 10:
                thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                processed_image = Image.fromarray(thresh)
                best_text = pytesseract.image_to_string(processed_image, lang='eng')
            
            # Normalize whitespace
            normalized_text = normalize_text(best_text) or best_text
            
            if not normalized_text or len(normalized_text.strip()) < 10:
                logger.warning(
                    f"OCR extracted minimal or no text from image: {filename}",
                    extra={"file_name": filename, "text_length": len(normalized_text)}
                )
            else:
                logger.info(
                    f"✅ Enhanced OCR extraction completed for {filename} (extracted {len(normalized_text.strip())} chars)",
                    extra={"file_name": filename, "text_length": len(normalized_text.strip())}
                )
            
            return normalized_text
            
        except Exception as e:
            logger.error(
                f"Error extracting text from image using OCR: {e}",
                extra={"error": str(e), "file_name": filename}
            )
            raise ValueError(f"Failed to extract text from image using OCR: {e}")
    
    def _extract_scanned_pdf_text(self, file_content: bytes, filename: str = "resume.pdf") -> str:
        """
        Extract text from scanned PDF using OCR.
        Converts PDF pages to images and then performs OCR on each page.
        
        Args:
            file_content: The binary content of the PDF file
            filename: Name of the file (for logging)
        
        Returns:
            Extracted text content as string
        
        Raises:
            ValueError: If OCR or pdf2image is not available or extraction fails
        """
        if not OCR_AVAILABLE:
            raise ValueError(
                "OCR libraries not available. Install required packages:\n"
                "pip install pytesseract pillow opencv-python\n"
                "Also install Tesseract OCR:\n"
                "  - Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki\n"
                "  - Linux: sudo apt-get install tesseract-ocr\n"
                "  - macOS: brew install tesseract"
            )
        
        # Check if we have at least one method to convert PDF to images
        if not PDF2IMAGE_AVAILABLE and not PYMUPDF_AVAILABLE:
            raise ValueError(
                "No PDF to image conversion library available. Install one of:\n"
                "  Option 1: pip install pdf2image (requires poppler)\n"
                "  Option 2: pip install PyMuPDF (no external dependencies)\n"
                "For poppler (if using pdf2image):\n"
                "  - Windows: Download from https://github.com/oschwartz10612/poppler-windows/releases\n"
                "  - Linux: sudo apt-get install poppler-utils\n"
                "  - macOS: brew install poppler"
            )
        
        try:
            # Create temporary file for PDF content
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(file_content)
                temp_pdf_path = temp_file.name
            
            try:
                logger.info(
                    f"Converting PDF pages to images for OCR: {filename}",
                    extra={"file_name": filename}
                )
                
                # Convert PDF pages to images (300 DPI for better OCR accuracy)
                # Try pdf2image first, fallback to PyMuPDF if poppler is not available
                pages = None
                try:
                    pages = convert_from_path(temp_pdf_path, dpi=300)
                except Exception as convert_error:
                    # If poppler is not available, try PyMuPDF as fallback
                    if PYMUPDF_AVAILABLE:
                        logger.info(
                            f"pdf2image failed (poppler not available), trying PyMuPDF fallback: {filename}",
                            extra={"file_name": filename, "error": str(convert_error)}
                        )
                        try:
                            pdf_doc = fitz.open(temp_pdf_path)
                            pages = []
                            for page_num in range(len(pdf_doc)):
                                page = pdf_doc[page_num]
                                # Render page to image at 300 DPI
                                mat = fitz.Matrix(300/72, 300/72)  # 300 DPI
                                pix = page.get_pixmap(matrix=mat)
                                # Convert to PIL Image
                                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                                pages.append(img)
                            pdf_doc.close()
                            logger.info(
                                f"✅ Successfully extracted {len(pages)} pages using PyMuPDF fallback: {filename}",
                                extra={"file_name": filename, "pages": len(pages)}
                            )
                        except Exception as pymupdf_error:
                            logger.error(
                                f"PyMuPDF fallback also failed: {pymupdf_error}",
                                extra={"file_name": filename, "error": str(pymupdf_error)}
                            )
                            raise ValueError(f"Failed to convert PDF to images for OCR. Poppler not available and PyMuPDF fallback failed: {pymupdf_error}")
                    else:
                        logger.error(
                            f"Failed to convert PDF to images: {convert_error}. "
                            f"Poppler is not installed and PyMuPDF is not available. "
                            f"Install poppler or install PyMuPDF: pip install PyMuPDF",
                            extra={"file_name": filename, "error": str(convert_error)}
                        )
                        raise ValueError(f"Failed to convert PDF to images for OCR: {convert_error}")
                
                if not pages or len(pages) == 0:
                    logger.warning(f"No pages extracted from PDF for OCR: {filename}")
                    return ""
                
                text_parts = []
                for i, page in enumerate(pages):
                    logger.debug(f"Processing PDF page {i+1}/{len(pages)} with enhanced OCR")
                    
                    # Convert PIL image to OpenCV format for preprocessing
                    img_array = np.array(page)
                    original_size = page.size
                    
                    # Scale up if image is too small (better OCR accuracy)
                    if page.size[0] < 1200 or page.size[1] < 1200:
                        scale_factor = max(1200 / page.size[0], 1200 / page.size[1])
                        new_size = (int(page.size[0] * scale_factor), int(page.size[1] * scale_factor))
                        page = page.resize(new_size, Image.Resampling.LANCZOS)
                        img_array = np.array(page)
                        logger.debug(f"Scaled PDF page {i+1} from {original_size} to {new_size} for better OCR")
                    
                    # Convert RGB to BGR for OpenCV
                    if len(img_array.shape) == 3:
                        img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                        # Convert to grayscale
                        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                    else:
                        gray = img_array
                    
                    # Noise removal
                    denoised = cv2.medianBlur(gray, 3)
                    
                    # Try multiple preprocessing methods and OCR configurations
                    best_text = ""
                    best_length = 0
                    
                    # Method 1: OTSU thresholding
                    try:
                        thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                        processed_image = Image.fromarray(thresh)
                        page_text = pytesseract.image_to_string(processed_image, lang='eng', config='--psm 6')
                        if len(page_text.strip()) > best_length:
                            best_text = page_text
                            best_length = len(page_text.strip())
                    except:
                        pass
                    
                    # Method 2: Adaptive thresholding
                    try:
                        adaptive = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
                        processed_image = Image.fromarray(adaptive)
                        page_text = pytesseract.image_to_string(processed_image, lang='eng', config='--psm 6')
                        if len(page_text.strip()) > best_length:
                            best_text = page_text
                            best_length = len(page_text.strip())
                    except:
                        pass
                    
                    # Method 3: Try different PSM modes
                    psm_modes = ['--psm 6', '--psm 11', '--psm 12', '--psm 3']
                    for psm in psm_modes:
                        try:
                            thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                            processed_image = Image.fromarray(thresh)
                            page_text = pytesseract.image_to_string(processed_image, lang='eng', config=psm)
                            if len(page_text.strip()) > best_length:
                                best_text = page_text
                                best_length = len(page_text.strip())
                        except:
                            continue
                    
                    # Use the best result
                    if best_text.strip():
                        text_parts.append(best_text)
                        logger.debug(f"Extracted {best_length} chars from page {i+1} using best OCR method")
                    else:
                        # Last resort: try basic OCR without preprocessing
                        try:
                            page_text = pytesseract.image_to_string(page, lang='eng')
                            if page_text.strip():
                                text_parts.append(page_text)
                        except Exception as ocr_page_error:
                            logger.warning(
                                f"OCR failed for page {i+1} of {filename}: {ocr_page_error}",
                                extra={"file_name": filename, "page": i+1, "error": str(ocr_page_error)}
                            )
                            # Continue with other pages
                            continue
                
                # Combine all pages
                raw_text = "\n".join(text_parts)
                normalized_text = normalize_text(raw_text) or raw_text
                
                if not normalized_text or len(normalized_text.strip()) == 0:
                    logger.warning(
                        f"⚠️ OCR extraction returned empty text for scanned PDF: {filename}",
                        extra={"file_name": filename, "pages": len(pages)}
                    )
                else:
                    logger.info(
                        f"✅ OCR extraction completed for scanned PDF: {filename} "
                        f"(extracted {len(normalized_text.strip())} chars from {len(pages)} page(s))",
                        extra={"file_name": filename, "pages": len(pages), "text_length": len(normalized_text.strip())}
                    )
                
                return normalized_text
                
            finally:
                # Clean up temp file
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)
                    
        except Exception as e:
            logger.error(
                f"Error extracting text from scanned PDF using OCR: {e}",
                extra={"error": str(e), "file_name": filename}
            )
            raise ValueError(f"Failed to extract text from scanned PDF using OCR: {e}")
    
    def _extract_html_text(self, file_content: bytes, filename: str = "resume.html") -> str:
        """
        Extract text from HTML file by parsing DOM with enhanced extraction.
        Removes HTML tags and extracts visible text content, with special handling
        for contact information sections. Includes fallback for plain text HTML.
        
        Args:
            file_content: The binary content of the HTML file
            filename: Name of the file (for logging)
        
        Returns:
            Extracted text content as string
        
        Raises:
            ValueError: If HTML parsing fails completely
        """
        # First, try to decode as plain text (fallback for simple HTML files)
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            html_content = None
            
            for encoding in encodings:
                try:
                    html_content = file_content.decode(encoding, errors='ignore')
                    if html_content:
                        break
                except:
                    continue
            
            if html_content is None:
                html_content = file_content.decode('utf-8', errors='ignore')
            
            # HTML-specific: Remove forwarding headers and metadata (only for HTML files)
            # This prevents extracting "Forwarded By" emails instead of candidate emails
            html_content_cleaned = html_content
            if filename.lower().endswith(('.html', '.htm')):
                # Strategy 1: Remove everything before "Personal Profile" or "Name:" (forwarding section)
                # Find where the actual resume content starts
                personal_profile_marker = re.search(r'(?i)(Personal\s+Profile|Name\s*:|\bRESUME\b)', html_content_cleaned)
                if personal_profile_marker:
                    # Keep only content from "Personal Profile" onwards
                    html_content_cleaned = html_content_cleaned[personal_profile_marker.start():]
                    logger.debug(f"Removed forwarding section before 'Personal Profile' in {filename}")
                
                # Strategy 2: Remove forwarding sections that contain non-candidate contact info
                forwarding_patterns = [
                    r'(?i)Forwarded\s+By:.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)To:\s*\[.*?\]\s*\n.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)From:.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)Resume\s+Link:.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)Comments:.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)I\s+thought\s+you\s+might\s+be\s+interested.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)This\s+resume\s+has\s+been\s+forwarded.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)This\s+email\s+was\s+sent.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)Email\s+ID\s+[A-Z0-9]+.*?(?=\n|Personal\s+Profile|Name\s*:)',
                    r'(?i)If\s+you\s+have\s+questions.*?CareerBuilder.*?(?=\n|Personal\s+Profile|Name\s*:)',
                ]
                
                for pattern in forwarding_patterns:
                    html_content_cleaned = re.sub(pattern, '', html_content_cleaned, flags=re.DOTALL)
                
                # Strategy 3: Remove lines that contain forwarding metadata
                lines = html_content_cleaned.split('\n')
                cleaned_lines = []
                in_forwarding_section = False
                found_personal_profile = False
                
                for i, line in enumerate(lines):
                    line_lower = line.lower().strip()
                    
                    # Detect start of forwarding section (only before Personal Profile)
                    if not found_personal_profile and any(keyword in line_lower for keyword in ['forwarded by', 'to:', 'from:', 'resume link:', 'comments:', 'i thought you might be interested', 'this resume has been forwarded', 'this email was sent', 'email id']):
                        in_forwarding_section = True
                        continue
                    
                    # Detect end of forwarding section (start of actual resume)
                    if any(marker in line_lower for marker in ['personal profile', 'name:', 'phone:', 'email:']):
                        in_forwarding_section = False
                        found_personal_profile = True
                    
                    # Skip lines in forwarding section
                    if in_forwarding_section:
                        continue
                    
                    # Skip separator lines that are part of forwarding headers (only before Personal Profile)
                    if not found_personal_profile and line.strip() == '' and i > 0 and i < len(lines) - 1:
                        prev_line = lines[i-1].lower().strip()
                        next_line = lines[i+1].lower().strip() if i+1 < len(lines) else ''
                        if any(keyword in prev_line for keyword in ['forwarded by', 'to:', 'from:']) or any(keyword in next_line for keyword in ['forwarded by', 'to:', 'from:']):
                            continue
                    
                    cleaned_lines.append(line)
                
                html_content_cleaned = '\n'.join(cleaned_lines)
                logger.debug(f"Cleaned HTML forwarding headers from {filename}, found Personal Profile: {found_personal_profile}")
            
            # If HTML parsing is available, use BeautifulSoup
            if HTML_PARSING_AVAILABLE:
                try:
                    # Try parsing with lxml first (faster)
                    try:
                        soup = BeautifulSoup(html_content_cleaned, 'lxml')
                    except:
                        # Fallback to html.parser if lxml fails
                        soup = BeautifulSoup(html_content_cleaned, 'html.parser')
                    
                    # Remove script and style elements (they don't contain visible text)
                    for script in soup(["script", "style", "meta", "link"]):
                        script.decompose()
                    
                    # Extract text from specific sections that often contain contact info
                    text_parts = []
                    
                    # Extract from header (often contains contact info)
                    header = soup.find('header')
                    if header:
                        header_text = header.get_text(separator=' ', strip=True)
                        if header_text:
                            text_parts.append(header_text)
                    
                    # Extract from title (may contain name)
                    title = soup.find('title')
                    if title:
                        title_text = title.get_text(strip=True)
                        if title_text:
                            text_parts.append(title_text)
                    
                    # Extract from <pre> tags (common in resume HTML files)
                    pre_tags = soup.find_all('pre')
                    for pre in pre_tags:
                        pre_text = pre.get_text(separator='\n', strip=True)
                        if pre_text:
                            text_parts.append(pre_text)
                    
                    # Extract from elements with contact-related classes/ids
                    contact_keywords = ['contact', 'email', 'phone', 'mobile', 'header', 'footer', 'info']
                    for keyword in contact_keywords:
                        # Find by class
                        elements = soup.find_all(class_=lambda x: x and keyword.lower() in str(x).lower())
                        for elem in elements:
                            elem_text = elem.get_text(separator=' ', strip=True)
                            if elem_text:
                                text_parts.append(elem_text)
                        
                        # Find by id
                        elem = soup.find(id=lambda x: x and keyword.lower() in str(x).lower())
                        if elem:
                            elem_text = elem.get_text(separator=' ', strip=True)
                            if elem_text:
                                text_parts.append(elem_text)
                    
                    # Extract all text content
                    all_text = soup.get_text(separator=' ', strip=True)
                    if all_text:
                        text_parts.append(all_text)
                    
                    # Try to extract text from embedded images using OCR (if available)
                    if OCR_AVAILABLE:
                        try:
                            # Find all img tags with src attributes
                            img_tags = soup.find_all('img')
                            for img in img_tags:
                                src = img.get('src', '')
                                if src:
                                    # Handle base64 encoded images
                                    if src.startswith('data:image'):
                                        try:
                                            # Extract base64 data
                                            header, encoded = src.split(',', 1)
                                            import base64
                                            image_data = base64.b64decode(encoded)
                                            # Perform OCR on the image
                                            image = Image.open(BytesIO(image_data))
                                            if image.mode != 'RGB':
                                                image = image.convert('RGB')
                                            # Pre-process for OCR
                                            img_array = np.array(image)
                                            img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                                            gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                                            thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                                            processed_image = Image.fromarray(thresh)
                                            ocr_text = pytesseract.image_to_string(processed_image, lang='eng')
                                            if ocr_text and ocr_text.strip():
                                                text_parts.append(ocr_text.strip())
                                                logger.debug(f"Extracted {len(ocr_text.strip())} chars from embedded image in HTML")
                                        except Exception as img_ocr_error:
                                            logger.debug(f"Failed to extract text from embedded image: {img_ocr_error}")
                        except Exception as ocr_error:
                            logger.debug(f"Failed to extract images from HTML for OCR: {ocr_error}")
                    
                    # Combine all text parts
                    combined_text = ' '.join(text_parts) if text_parts else all_text
                    
                    # Normalize whitespace
                    normalized_text = normalize_text(combined_text) if combined_text else ""
                    
                    if normalized_text and len(normalized_text.strip()) > 10:
                        logger.info(
                            f"✅ Enhanced HTML text extraction completed for {filename} (extracted {len(normalized_text.strip())} chars)",
                            extra={"file_name": filename, "text_length": len(normalized_text.strip())}
                        )
                        return normalized_text
                    
                except Exception as bs_error:
                    logger.warning(
                        f"BeautifulSoup parsing failed for {filename}, trying plain text fallback: {bs_error}",
                        extra={"file_name": filename, "error": str(bs_error)}
                    )
            
            # Fallback: Extract text directly using regex (for simple HTML or when BeautifulSoup fails)
            # Use cleaned content (with forwarding headers removed)
            text = re.sub(r'<[^>]+>', ' ', html_content_cleaned)
            # Decode HTML entities
            try:
                text = html.unescape(text)
            except:
                pass
            # Clean up whitespace
            text = ' '.join(text.split())
            
            if text and len(text.strip()) > 10:
                normalized_text = normalize_text(text) or text
                logger.info(
                    f"✅ HTML text extraction completed using fallback method for {filename} (extracted {len(normalized_text.strip())} chars)",
                    extra={"file_name": filename, "text_length": len(normalized_text.strip()), "method": "regex_fallback"}
                )
                return normalized_text
            
            # Last resort: return decoded content as-is (use cleaned content)
            if html_content_cleaned and len(html_content_cleaned.strip()) > 10:
                normalized_text = normalize_text(html_content_cleaned) or html_content_cleaned
                logger.info(
                    f"✅ HTML text extraction completed using raw decode for {filename} (extracted {len(normalized_text.strip())} chars)",
                    extra={"file_name": filename, "text_length": len(normalized_text.strip()), "method": "raw_decode"}
                )
                return normalized_text
            
            # If no text found, try OCR fallback for HTML files (in case it's an image-based HTML)
            if OCR_AVAILABLE and (PYMUPDF_AVAILABLE or PDF2IMAGE_AVAILABLE):
                try:
                    logger.info(
                        f"🔄 HTML extraction returned insufficient text, trying OCR fallback for {filename}",
                        extra={"file_name": filename, "fallback_type": "html_ocr"}
                    )
                    # Try to convert HTML to image and run OCR
                    # This is useful for HTML files that are actually screenshots or image-based
                    # For now, we'll try to extract any embedded images and run OCR on them
                    # If that doesn't work, we'll return the minimal text we have
                    # Note: Full HTML-to-image conversion would require additional libraries like playwright/selenium
                except Exception as ocr_fallback_error:
                    logger.debug(f"OCR fallback for HTML failed: {ocr_fallback_error}")
            
            raise ValueError("No extractable text found in HTML file")
            
        except Exception as e:
            logger.error(
                f"Error extracting text from HTML: {e}",
                extra={"error": str(e), "file_name": filename}
            )
            # Last resort: try to decode as plain text
            try:
                text = file_content.decode('utf-8', errors='ignore')
                if text and len(text.strip()) > 10:
                    logger.warning(f"Using raw UTF-8 decode as last resort for {filename}")
                    return normalize_text(text) or text
            except:
                pass
            raise ValueError(f"Failed to extract text from HTML: {e}")
    
    async def extract_text_with_fallback(self, file_content: bytes, filename: str, original_text: str = None) -> str:
        """
        Extract text with enhanced fallback methods for image/HTML/DOCX/PDF files.
        Used when initial extraction fails or returns insufficient/null results.
        
        Args:
            file_content: The binary content of the file
            filename: Name of the file
            original_text: Previously extracted text (if any)
        
        Returns:
            Extracted text content as string
        """
        filename_lower = filename.lower()
        
        # If it's an image file, try enhanced OCR
        if filename_lower.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif')):
            logger.info(
                f"🔄 FALLBACK: Retrying with enhanced OCR for image file: {filename}",
                extra={"file_name": filename, "fallback_type": "enhanced_ocr"}
            )
            try:
                enhanced_text = self._extract_image_text(file_content, filename)
                if enhanced_text and len(enhanced_text.strip()) > len(original_text.strip() if original_text else ""):
                    logger.info(
                        f"✅ FALLBACK SUCCESS: Enhanced OCR extracted {len(enhanced_text.strip())} chars from {filename}",
                        extra={"file_name": filename, "text_length": len(enhanced_text.strip())}
                    )
                    return enhanced_text
            except Exception as e:
                logger.warning(f"Enhanced OCR fallback failed: {e}", extra={"file_name": filename, "error": str(e)})
        
        # If it's an HTML file, try enhanced HTML parsing and OCR on embedded images
        elif filename_lower.endswith(('.html', '.htm')):
            logger.info(
                f"🔄 FALLBACK: Retrying with enhanced HTML parsing and OCR for: {filename}",
                extra={"file_name": filename, "fallback_type": "enhanced_html_ocr"}
            )
            try:
                # First try enhanced HTML parsing
                enhanced_text = self._extract_html_text(file_content, filename)
                
                # If HTML parsing didn't help much, try OCR on embedded images
                if (not enhanced_text or len(enhanced_text.strip()) < 100) and OCR_AVAILABLE:
                    try:
                        # Decode HTML to find embedded images
                        html_content = file_content.decode('utf-8', errors='ignore')
                        if HTML_PARSING_AVAILABLE:
                            soup = BeautifulSoup(html_content, 'html.parser')
                            img_tags = soup.find_all('img')
                            ocr_text_parts = []
                            
                            for img in img_tags:
                                src = img.get('src', '')
                                if src and src.startswith('data:image'):
                                    try:
                                        # Extract base64 image data
                                        header, encoded = src.split(',', 1)
                                        import base64
                                        image_data = base64.b64decode(encoded)
                                        # Perform OCR
                                        image = Image.open(BytesIO(image_data))
                                        if image.mode != 'RGB':
                                            image = image.convert('RGB')
                                        # Pre-process for OCR
                                        img_array = np.array(image)
                                        img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                                        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                                        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                                        processed_image = Image.fromarray(thresh)
                                        ocr_text = pytesseract.image_to_string(processed_image, lang='eng')
                                        if ocr_text and ocr_text.strip():
                                            ocr_text_parts.append(ocr_text.strip())
                                    except Exception as img_ocr_error:
                                        logger.debug(f"Failed to OCR embedded image in HTML: {img_ocr_error}")
                            
                            if ocr_text_parts:
                                ocr_combined = ' '.join(ocr_text_parts)
                                if len(ocr_combined.strip()) > len(enhanced_text.strip() if enhanced_text else ""):
                                    enhanced_text = ocr_combined
                                    logger.info(
                                        f"✅ OCR on HTML embedded images extracted {len(enhanced_text.strip())} chars",
                                        extra={"file_name": filename, "text_length": len(enhanced_text.strip())}
                                    )
                    except Exception as html_ocr_error:
                        logger.debug(f"HTML OCR fallback failed: {html_ocr_error}")
                
                if enhanced_text and len(enhanced_text.strip()) > len(original_text.strip() if original_text else ""):
                    logger.info(
                        f"✅ FALLBACK SUCCESS: Enhanced HTML parsing/OCR extracted {len(enhanced_text.strip())} chars from {filename}",
                        extra={"file_name": filename, "text_length": len(enhanced_text.strip())}
                    )
                    return enhanced_text
            except Exception as e:
                logger.warning(f"Enhanced HTML parsing fallback failed: {e}", extra={"file_name": filename, "error": str(e)})
        
        # If it's a DOCX file, try enhanced extraction with XML fallback
        elif filename_lower.endswith('.docx'):
            logger.info(
                f"🔄 FALLBACK: Retrying with enhanced DOCX extraction for: {filename}",
                extra={"file_name": filename, "fallback_type": "enhanced_docx"}
            )
            try:
                enhanced_text = self._extract_docx_text(file_content, filename)
                if enhanced_text and len(enhanced_text.strip()) > len(original_text.strip() if original_text else ""):
                    logger.info(
                        f"✅ FALLBACK SUCCESS: Enhanced DOCX extraction extracted {len(enhanced_text.strip())} chars from {filename}",
                        extra={"file_name": filename, "text_length": len(enhanced_text.strip())}
                    )
                    return enhanced_text
            except Exception as e:
                logger.warning(f"Enhanced DOCX extraction fallback failed: {e}", extra={"file_name": filename, "error": str(e)})
        
        # If it's a PDF that might be scanned, try OCR
        elif filename_lower.endswith('.pdf'):
            logger.info(
                f"🔄 FALLBACK: Retrying with OCR for potentially scanned PDF: {filename}",
                extra={"file_name": filename, "fallback_type": "pdf_ocr"}
            )
            try:
                if OCR_AVAILABLE and PDF2IMAGE_AVAILABLE:
                    enhanced_text = self._extract_scanned_pdf_text(file_content, filename)
                    if enhanced_text and len(enhanced_text.strip()) > len(original_text.strip() if original_text else ""):
                        logger.info(
                            f"✅ FALLBACK SUCCESS: OCR extracted {len(enhanced_text.strip())} chars from scanned PDF {filename}",
                            extra={"file_name": filename, "text_length": len(enhanced_text.strip())}
                        )
                        return enhanced_text
            except Exception as e:
                logger.warning(f"PDF OCR fallback failed: {e}", extra={"file_name": filename, "error": str(e)})
        
        # Return original text if fallback didn't improve it
        return original_text if original_text else ""

