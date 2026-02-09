"""
Resume Designation Extractor for Google Colab
A single-file solution for extracting job designation/role from resumes using Ollama LLM.

USAGE INSTRUCTIONS FOR GOOGLE COLAB:
====================================

1. Install required dependencies:
   !pip install httpx PyPDF2 python-docx ollama

2. Install and start Ollama in Colab:
   !curl -fsSL https://ollama.com/install.sh | sh
   !ollama serve &
   !sleep 5  # Wait for Ollama to start
   
3. Pull the required model (llama3.1):
   !ollama pull llama3.1

4. Upload this ResumeParser.py file to your Colab environment

5. Run the parser:
   import asyncio
   from ResumeParser import main
   result = asyncio.run(main())

   OR use directly:
   import asyncio
   from ResumeParser import parse_resume_file
   result = asyncio.run(parse_resume_file())

6. The script will prompt you to upload a resume file (PDF, DOCX, or TXT)
   and will display the parsed results.

CUSTOMIZATION:
- To use a different Ollama host: result = asyncio.run(parse_resume_file(ollama_host="http://your-host:11434"))
- To use a different model: result = asyncio.run(parse_resume_file(model_name="your-model"))

OUTPUT:
The function returns a dictionary with the following key:
- designation: The candidate's job title/role (exactly as written in the resume) or null if not found

The designation is extracted following strict rules:
1. Current/present role (if explicitly marked)
2. Most recent experience entry
3. Resume headline or summary
4. First occurrence if multiple at same level
"""

import json
import re
import gc
from typing import Dict, Optional, List
from io import BytesIO
import httpx
from httpx import Timeout

# Try to import Google Colab file upload
try:
    from google.colab import files
    COLAB_AVAILABLE = True
except ImportError:
    COLAB_AVAILABLE = False
    print("Note: Google Colab not detected. Using standard file input.")

# Try to import OLLAMA Python client
try:
    import ollama
    OLLAMA_CLIENT_AVAILABLE = True
except ImportError:
    OLLAMA_CLIENT_AVAILABLE = False
    print("Note: OLLAMA Python client not available, using HTTP API directly")

# PDF and DOCX parsing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not available. PDF files cannot be processed.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX files cannot be processed.")

# Try to import textract for .doc file support
try:
    import textract
    TEXTTRACT_AVAILABLE = True
except ImportError:
    TEXTTRACT_AVAILABLE = False
    print("Note: textract not available. .doc files will use fallback methods.")

# Try to import Apache Tika for .doc file support
try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    print("Note: Apache Tika not available. Will use alternative methods for .doc files.")

# Try to import olefile for .doc file support (alternative method)
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    print("Note: olefile not available. Will use alternative methods for .doc files.")

# Check for antiword command-line tool
import shutil
ANTIWORD_AVAILABLE = shutil.which("antiword") is not None
if not ANTIWORD_AVAILABLE:
    print("Note: antiword command not found. Install it for better .doc extraction.")

# Check for LibreOffice command-line tool
LIBREOFFICE_AVAILABLE = shutil.which("soffice") is not None or shutil.which("libreoffice") is not None
if not LIBREOFFICE_AVAILABLE:
    print("Note: LibreOffice not found. Install it for most reliable .doc conversion.")


# Configuration
OLLAMA_HOST = "http://localhost:11434"  # Default Ollama host
MODEL_NAME = "llama3.1"  # Default model

MASTER_PROMPT = """
ROLE:
You are an ATS resume parsing expert specializing in US IT staffing profiles.

CONTEXT:
Candidate profiles and resumes may be unstructured and inconsistently formatted.
Designation refers to the candidate's explicitly stated current or most recent job title.

TASK:
Extract the candidate's designation (job title) from the profile text.

SELECTION RULES (IN ORDER OF PRIORITY):
1. If a title is explicitly marked as "current", "present", or equivalent, select that.
2. Else, select the title associated with the most recent experience entry.
3. Else, select the designation mentioned in the resume headline or summary.
4. If multiple titles appear at the same level, select the first occurrence.

CONSTRAINTS:
- Extract only one designation.
- Preserve the designation exactly as written.
- Do not infer or normalize titles.
- Do not include company names, skills, durations, or locations.
- Ignore aspirational, desired, or target roles.

ANTI-HALLUCINATION RULES:
- If no explicit designation is found, return null.
- Never guess or infer a designation.
- Do not derive designation from skills, certifications, or projects.

OUTPUT FORMAT:
Return only valid JSON. No additional text.

JSON SCHEMA:
{
  "designation": "string | null"
}
"""


# Utility Functions
def normalize_phone(phone: Optional[str]) -> Optional[str]:
    """Normalize phone number to E.164 format if possible."""
    if not phone:
        return None
    
    cleaned = re.sub(r'[^\d+]', '', phone.strip())
    
    if cleaned.startswith('+'):
        if len(cleaned) >= 10:
            return cleaned
    else:
        if len(cleaned) == 10:
            return f"+1{cleaned}"
        elif len(cleaned) == 11 and cleaned.startswith('1'):
            return f"+{cleaned}"
    
    return cleaned if cleaned else None


def normalize_email(email: Optional[str]) -> Optional[str]:
    """Normalize email to lowercase."""
    if not email:
        return None
    
    email = email.strip().lower()
    email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    if re.match(email_pattern, email):
        return email
    return None


def extract_skills(text: Optional[str]) -> List[str]:
    """Extract and normalize skills from text."""
    if not text:
        return []
    
    separators = r'[,;|•\n\r\t]+'
    raw_skills = re.split(separators, text)
    
    skills = []
    for skill in raw_skills:
        skill = skill.strip()
        if skill and len(skill) > 1:
            skill = re.sub(r'\s+', ' ', skill)
            skill = skill.title()
            skills.append(skill)
    
    return skills[:50]


def normalize_text(text: Optional[str]) -> Optional[str]:
    """Normalize text by removing extra whitespace."""
    if not text:
        return None
    text = re.sub(r'\s+', ' ', text.strip())
    return text if text else None


def extract_pdf_text(file_content: bytes) -> str:
    """Extract text from PDF file."""
    if not PDF_AVAILABLE:
        raise ValueError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
    
    pdf_file = BytesIO(file_content)
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []
        for page in pdf_reader.pages:
            text_parts.append(page.extract_text())
        text = "\n".join(text_parts)
        # Explicitly clean up reader
        del pdf_reader
        return text
    finally:
        # Explicitly close BytesIO
        pdf_file.close()
        del pdf_file


def extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx is not installed. Install it with: pip install python-docx")
    
    doc_file = BytesIO(file_content)
    try:
        doc = Document(doc_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        text = "\n".join(text_parts)
        # Explicitly clean up document object
        del doc
        return text
    finally:
        # Explicitly close BytesIO
        doc_file.close()
        del doc_file


def extract_doc_text(file_content: bytes) -> str:
    """
    Extract text from DOC file (older Microsoft Word format).
    Uses multiple methods in order of reliability:
    1. LibreOffice headless conversion (most reliable for production)
    2. antiword (good for plain-text extraction)
    3. Apache Tika (read text with layout)
    4. textract (if available)
    5. python-docx fallback (might work for some files)
    6. olefile (basic binary extraction)
    """
    import os
    import subprocess
    import tempfile
    
    # Create temporary file for .doc content
    with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
        temp_file.write(file_content)
        temp_doc_path = temp_file.name
    
    try:
        # Method 1: LibreOffice headless conversion (MOST RELIABLE FOR PRODUCTION)
        if LIBREOFFICE_AVAILABLE:
            try:
                print("Attempting .doc extraction using LibreOffice headless conversion")
                # Create temp directory for output
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Use LibreOffice to convert .doc to .docx
                    libreoffice_cmd = shutil.which("soffice") or shutil.which("libreoffice")
                    cmd = [
                        libreoffice_cmd,
                        "--headless",
                        "--convert-to", "docx",
                        "--outdir", temp_dir,
                        temp_doc_path
                    ]
                    
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    if result.returncode == 0:
                        # Find the converted .docx file
                        converted_docx = os.path.join(temp_dir, os.path.basename(temp_doc_path).replace('.doc', '.docx'))
                        if os.path.exists(converted_docx):
                            # Read the converted .docx and extract text
                            with open(converted_docx, 'rb') as f:
                                docx_content = f.read()
                            text = extract_docx_text(docx_content)
                            # Free docx_content immediately after extraction
                            del docx_content
                            gc.collect()
                            if text.strip():
                                print("Successfully extracted .doc file using LibreOffice conversion")
                                return text
            except subprocess.TimeoutExpired:
                print("LibreOffice conversion timed out")
            except Exception as lo_error:
                print(f"LibreOffice conversion failed: {lo_error}")
        
        # Method 2: antiword (GOOD FOR PLAIN-TEXT EXTRACTION)
        if ANTIWORD_AVAILABLE:
            try:
                print("Attempting .doc extraction using antiword")
                result = subprocess.run(
                    ["antiword", temp_doc_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    print("Successfully extracted .doc file using antiword")
                    return result.stdout
            except subprocess.TimeoutExpired:
                print("antiword extraction timed out")
            except Exception as aw_error:
                print(f"antiword extraction failed: {aw_error}")
        
        # Method 3: Apache Tika (READ TEXT WITH LAYOUT)
        if TIKA_AVAILABLE:
            try:
                print("Attempting .doc extraction using Apache Tika")
                parsed = tika_parser.from_file(temp_doc_path)
                if parsed and 'content' in parsed and parsed['content']:
                    text = parsed['content'].strip()
                    if text:
                        print("Successfully extracted .doc file using Apache Tika")
                        return text
            except Exception as tika_error:
                print(f"Apache Tika extraction failed: {tika_error}")
        
        # Method 4: textract (if available)
        if TEXTTRACT_AVAILABLE:
            try:
                print("Attempting .doc extraction using textract")
                extracted_bytes = textract.process(temp_doc_path)
                text = extracted_bytes.decode('utf-8', errors='ignore')
                if text.strip():
                    print("Successfully extracted .doc file using textract")
                    return text
            except Exception as textract_error:
                print(f"textract extraction failed: {textract_error}")
        
        # Method 5: Try python-docx as fallback (might work for some .doc files that are actually .docx)
        if DOCX_AVAILABLE:
            try:
                print("Attempting .doc extraction using python-docx fallback")
                doc_file = BytesIO(file_content)
                try:
                    doc = Document(doc_file)
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        text_parts.append(paragraph.text)
                    # Also try to extract from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                row_text.append(cell.text)
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                    extracted_text = "\n".join(text_parts)
                    if extracted_text.strip():
                        print("Successfully extracted .doc file using python-docx fallback")
                        del doc
                        return extracted_text
                    del doc
                finally:
                    doc_file.close()
                    del doc_file
            except Exception:
                pass
        
        # Method 6: Try to extract using olefile (for binary .doc files)
        if OLEFILE_AVAILABLE:
            ole = None
            doc_file = None
            try:
                print("Attempting .doc extraction using olefile")
                # .doc files are OLE compound documents
                doc_file = BytesIO(file_content)
                ole = olefile.OleFileIO(doc_file)
                # Try to find WordDocument stream
                if ole.exists('WordDocument'):
                    stream = ole.openstream('WordDocument')
                    # Read and try to extract text (basic extraction)
                    data = stream.read()
                    stream.close()
                    # Simple text extraction from binary data
                    # This is a basic approach - look for readable text
                    text_chunks = []
                    current_chunk = b""
                    for byte in data:
                        if 32 <= byte <= 126 or byte in [9, 10, 13]:  # Printable ASCII
                            current_chunk += bytes([byte])
                        else:
                            if len(current_chunk) > 3:
                                try:
                                    text_chunks.append(current_chunk.decode('ascii', errors='ignore'))
                                except:
                                    pass
                            current_chunk = b""
                    if current_chunk:
                        try:
                            text_chunks.append(current_chunk.decode('ascii', errors='ignore'))
                        except:
                            pass
                    del data  # Free memory immediately
                    extracted_text = "\n".join(text_chunks)
                    del text_chunks  # Free memory immediately
                    if extracted_text.strip():
                        print("Successfully extracted .doc file using olefile")
                        return extracted_text
            except Exception as ole_error:
                print(f"olefile extraction failed: {ole_error}")
            finally:
                # Always close olefile and BytesIO
                if ole is not None:
                    try:
                        ole.close()
                    except:
                        pass
                    del ole
                if doc_file is not None:
                    try:
                        doc_file.close()
                    except:
                        pass
                    del doc_file
        
        # If all methods fail, raise an error with helpful message
        raise ValueError(
            "Failed to extract text from .doc file using all available methods.\n"
            "Installation options (in order of reliability):\n"
            "1. LibreOffice (headless): Most reliable for production\n"
            "   - Windows: Download from https://www.libreoffice.org/\n"
            "   - Linux: sudo apt-get install libreoffice\n"
            "2. antiword: Good for plain-text extraction\n"
            "   - Windows: Download from http://www.winfield.demon.nl/\n"
            "   - Linux: sudo apt-get install antiword\n"
            "3. Apache Tika: pip install tika (requires Java)\n"
            "4. textract: pip install textract (may require system dependencies)\n"
            "5. olefile: pip install olefile (basic support, already installed)\n"
            "6. Convert .doc files to .docx format before processing"
        )
    finally:
        # Clean up temp file
        if os.path.exists(temp_doc_path):
            os.unlink(temp_doc_path)


def extract_text(file_content: bytes, filename: str) -> str:
    """Extract text from uploaded file based on extension."""
    try:
        if filename.lower().endswith('.pdf'):
            return extract_pdf_text(file_content)
        elif filename.lower().endswith('.docx'):
            return extract_docx_text(file_content)
        elif filename.lower().endswith('.doc'):
            return extract_doc_text(file_content)
        elif filename.lower().endswith('.txt'):
            return file_content.decode('utf-8', errors='ignore')
        else:
            return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        raise ValueError(f"Failed to extract text from file: {e}")


async def check_ollama_connection(ollama_host: str = OLLAMA_HOST) -> tuple[bool, Optional[str]]:
    """Check if OLLAMA is accessible and running. Returns (is_connected, available_model)."""
    try:
        async with httpx.AsyncClient(timeout=Timeout(5.0)) as client:
            response = await client.get(f"{ollama_host}/api/tags")
            if response.status_code == 200:
                models_data = response.json()
                models = models_data.get("models", [])
                for model in models:
                    model_name = model.get("name", "")
                    if "llama3.1" in model_name.lower() or "llama3" in model_name.lower():
                        return True, model_name
                if models:
                    return True, models[0].get("name", "")
                return True, None
            return False, None
    except Exception as e:
        print(f"Failed to check OLLAMA connection: {e}")
        return False, None


def extract_json(text: str) -> Dict:
    """Extract JSON object from LLM response."""
    json_match = re.search(r'\{.*\}', text, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError:
            pass
    
    try:
        return json.loads(text.strip())
    except json.JSONDecodeError:
        print(f"Warning: Failed to parse JSON from LLM response. Response preview: {text[:500]}")
        return {
            "designation": None,
        }


async def parse_resume_with_ollama(
    resume_text: str,
    filename: str,
    ollama_host: str = OLLAMA_HOST,
    model_name: str = MODEL_NAME
) -> Dict:
    """Parse resume text using Ollama LLM and return structured data."""
    import asyncio
    
    # Check OLLAMA connection first
    is_connected, available_model = await check_ollama_connection(ollama_host)
    if not is_connected:
        raise RuntimeError(
            f"OLLAMA is not accessible at {ollama_host}. "
            "Please ensure OLLAMA is running. In Google Colab, you may need to install and start Ollama."
        )
    
    # Use available model if llama3.1 not found
    model_to_use = model_name
    if available_model and "llama3.1" not in available_model.lower():
        print(f"Warning: llama3.1 not found, using available model: {available_model}")
        model_to_use = available_model
    
    # Prepare prompt
    prompt = f"{MASTER_PROMPT}\n\nInput: {resume_text[:10000]}\n\nOutput:"
    
    result = None
    last_error = None
    
    # Try using OLLAMA Python client first
    if OLLAMA_CLIENT_AVAILABLE:
        try:
            loop = asyncio.get_event_loop()
            def _generate():
                client = ollama.Client(host=ollama_host.replace("http://", "").replace("https://", ""))
                response = client.generate(
                    model=model_to_use,
                    prompt=prompt,
                    options={
                        "temperature": 0.1,
                        "top_p": 0.9,
                    }
                )
                return {"response": response.get("response", "")}
            
            result = await loop.run_in_executor(None, _generate)
            print("Successfully used OLLAMA Python client")
        except Exception as e:
            print(f"OLLAMA Python client failed, falling back to HTTP API: {e}")
            result = None
    
    # Fallback to HTTP API
    if result is None:
        async with httpx.AsyncClient(timeout=Timeout(1200.0)) as client:
            # Try /api/generate endpoint
            try:
                response = await client.post(
                    f"{ollama_host}/api/generate",
                    json={
                        "model": model_to_use,
                        "prompt": prompt,
                        "stream": False,
                        "options": {
                            "temperature": 0.1,
                            "top_p": 0.9,
                        }
                    }
                )
                response.raise_for_status()
                result = response.json()
                print("Successfully used /api/generate endpoint")
            except httpx.HTTPStatusError as e:
                if e.response.status_code == 404:
                    # Try /api/chat endpoint
                    print("OLLAMA /api/generate not found, trying /api/chat endpoint")
                    try:
                        response = await client.post(
                            f"{ollama_host}/api/chat",
                            json={
                                "model": model_to_use,
                                "messages": [
                                    {"role": "user", "content": prompt}
                                ],
                                "stream": False,
                                "options": {
                                    "temperature": 0.1,
                                    "top_p": 0.9,
                                }
                            }
                        )
                        response.raise_for_status()
                        result = response.json()
                        if "message" in result and "content" in result["message"]:
                            result = {"response": result["message"]["content"]}
                        else:
                            raise ValueError("Unexpected response format from OLLAMA chat API")
                        print("Successfully used /api/chat endpoint")
                    except Exception as e2:
                        last_error = e2
                        print(f"OLLAMA /api/chat also failed: {e2}")
                else:
                    raise
    
    if result is None:
        raise RuntimeError(
            f"All OLLAMA API endpoints failed. "
            f"OLLAMA is running at {ollama_host} but endpoints return errors. "
            f"Last error: {last_error}"
        )
    
    # Extract JSON from response
    raw_output = result.get("response", "")
    parsed_data = extract_json(raw_output)
    
    # Ensure designation field exists (even if null)
    if "designation" not in parsed_data:
        parsed_data["designation"] = None
    
    # Preserve designation exactly as written (no normalization per requirements)
    # Just ensure it's a string or null
    if parsed_data["designation"] is not None:
        parsed_data["designation"] = str(parsed_data["designation"]).strip()
        if not parsed_data["designation"]:
            parsed_data["designation"] = None
    
    return parsed_data


def upload_file_colab():
    """Upload file in Google Colab environment."""
    if not COLAB_AVAILABLE:
        raise RuntimeError(
            "Google Colab file upload not available. "
            "Please run this in Google Colab or provide file content directly."
        )
    
    uploaded = files.upload()
    if not uploaded:
        raise ValueError("No file was uploaded")
    
    # Get the first uploaded file
    filename = list(uploaded.keys())[0]
    file_content = uploaded[filename]
    
    return filename, file_content


async def parse_resume_file(
    file_content: bytes = None,
    filename: str = None,
    ollama_host: str = OLLAMA_HOST,
    model_name: str = MODEL_NAME
) -> Dict:
    """
    Main function to parse a resume file.
    
    Args:
        file_content: Bytes content of the file (optional if running in Colab)
        filename: Name of the file (optional if running in Colab)
        ollama_host: Ollama server host (default: http://localhost:11434)
        model_name: Ollama model name (default: llama3.1)
    
    Returns:
        Dictionary with parsed resume data
    """
    # Upload file if in Colab and no file_content provided
    if file_content is None:
        if COLAB_AVAILABLE:
            print("Uploading file from Google Colab...")
            filename, file_content = upload_file_colab()
        else:
            raise ValueError(
                "No file content provided. "
                "In Google Colab, the file will be uploaded automatically. "
                "Otherwise, provide file_content and filename parameters."
            )
    
    if not filename:
        filename = "resume.pdf"
    
    print(f"Processing file: {filename}")
    print(f"File size: {len(file_content)} bytes")
    
    # Extract text from file
    print("Extracting text from file...")
    resume_text = extract_text(file_content, filename)
    
    # Free file_content memory immediately after extraction
    del file_content
    gc.collect()  # Force garbage collection to free memory
    
    if not resume_text or len(resume_text.strip()) < 50:
        raise ValueError("Could not extract sufficient text from resume")
    
    print(f"Extracted {len(resume_text)} characters of text")
    print(f"Text preview (first 200 chars): {resume_text[:200]}...")
    
    # Parse resume using Ollama
    print(f"\nParsing resume with Ollama (model: {model_name})...")
    print(f"Ollama host: {ollama_host}")
    
    parsed_data = await parse_resume_with_ollama(
        resume_text,
        filename,
        ollama_host,
        model_name
    )
    
    # Free resume_text memory after parsing
    del resume_text
    gc.collect()  # Force garbage collection to free memory
    
    print("\n" + "="*50)
    print("DESIGNATION EXTRACTION COMPLETE")
    print("="*50)
    
    return parsed_data


# Main execution function for Colab
async def main():
    """Main function to run in Google Colab."""
    print("="*50)
    print("RESUME DESIGNATION EXTRACTOR - Google Colab Edition")
    print("="*50)
    print("\nThis script will:")
    print("1. Upload a resume file (PDF, DOCX, or TXT)")
    print("2. Extract text from the file")
    print("3. Extract the candidate's designation (job title) using Ollama LLM")
    print("4. Display the extracted designation")
    print("\n" + "="*50 + "\n")
    
    try:
        # Parse resume
        result = await parse_resume_file()
        
        # Display results
        print("\nExtracted Designation:")
        print(json.dumps(result, indent=2, ensure_ascii=False))
        if result and result.get("designation"):
            print(f"\n✅ Designation found: {result['designation']}")
        else:
            print("\n⚠️  No designation found in the resume.")
        
        return result
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return None


def run_parser(ollama_host: str = OLLAMA_HOST, model_name: str = MODEL_NAME):
    """
    Convenience function to run the parser synchronously.
    Use this in Google Colab notebooks for easier execution.
    
    Example:
        from ResumeParser import run_parser
        result = run_parser()
    """
    import asyncio
    return asyncio.run(parse_resume_file(ollama_host=ollama_host, model_name=model_name))


# For running in Colab with asyncio
if __name__ == "__main__":
    import asyncio
    
    # Run the main function
    result = asyncio.run(main())
    
    if result:
        print("\n✅ Resume parsing completed successfully!")
    else:
        print("\n❌ Resume parsing failed. Please check the error messages above.")

